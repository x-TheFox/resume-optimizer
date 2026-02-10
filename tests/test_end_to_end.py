
import unittest
import os
import json
from unittest.mock import patch, MagicMock

# Add src to python path if needed, but relative imports should work if structured correctly
# Assuming running from project root
import sys
sys.path.append(os.getcwd())

from src.research_engine import ResearchEngine
from src.llm_analyzer import LLMAnalyzer
from src.resume_editor import ResumeEditor
from src.pdf_generator import PDFGenerator
from docx import Document

class TestResumeOptimizer(unittest.TestCase):
    
    @classmethod
    def setUpClass(cls):
        # Config
        cls.test_docx_path = "tests/test_resume.docx"
        cls.output_dir = "tests/test_outputs"
        os.makedirs(cls.output_dir, exist_ok=True)
        
        # Test Data
        cls.job_title = "Senior Python Engineer"
        cls.job_description = "We are looking for a Senior Python Engineer to build scalable APIs."
        cls.company_name = "TechCorp"
        
        # Verify test file exists
        if not os.path.exists(cls.test_docx_path):
            raise FileNotFoundError(f"Test file not found: {cls.test_docx_path}")

    def test_1_research_engine_mock(self):
        """Test ResearchEngine with mocked search results."""
        print("\n--- Testing Research Engine (Mocked) ---")
        
        re = ResearchEngine()
        
        # Mock the internal search methods to avoid actual network calls and rate limits
        with patch.object(re, '_safe_search') as mock_search:
            mock_search.side_effect = [
                "Responsibility: Write clean Python code.", # role
                "Trend: Microservices and Kubernetes.",    # tech
                "Value: Innovation and integrity.",        # values
                "News: Launched new AI product.",          # news
                "Competitor: RivalCorp.",                  # competitors
                "Skill: FastAPI, Docker.",                 # shadow skills
            ]
            
            profile = re.research(self.job_title, self.job_description, self.company_name)
            
            self.assertEqual(profile['job_title'], self.job_title)
            self.assertIn("Python", profile['role_responsibilities'])
            self.assertEqual(profile['company_name'], self.company_name)
            print("Research Engine Profile Built Successfully.")

    def test_2_resume_parsing(self):
        """Test parsing of the .docx file."""
        print("\n--- Testing Document Parsing ---")
        editor = ResumeEditor(self.test_docx_path)
        text = editor.extract_text()
        self.assertIn("John Doe", text)
        self.assertIn("Backend Engineer", text)
        print("Resume text extracted successfully.")
        return text

    def test_3_llm_analyzer_mock(self):
        """Test LLM Analyzer with mocked Groq responses."""
        print("\n--- Testing LLM Analyzer (Mocked) ---")
        
        # Fake successes profile for testing
        success_profile = {
            "role_responsibilities": "Write code",
            "tech_trends": "AI",
            "company_values": "Kindness",
            "recent_news": "None",
            "competitors": "None",
            "shadow_skills": "None",
            "cultural_tone": "balanced"
        }
        
        analyzer = LLMAnalyzer(api_key="fake-key", model="fake-model")
        
        # Define expected responses mapped to prompts or just sequential with debug
        responses_iter = iter([
            # 1. Gap Analysis
            "The candidate has strong Python skills but lacks Kubernetes knowledge.",
            
            # 2. Section Scores
            json.dumps({"skills": 85, "experience": 90, "impact": 75}),
            
            # 3. Match Scores
            json.dumps({"technical_match": 88, "cultural_match": 80}),
            
            # 4. ATS Sim
            json.dumps({"score": 92, "warnings": ["Avoid columns"]}),
            
            # 5. Suggestions
            json.dumps([
                {
                    "section": "Experience",
                    "original_text": "microservices architecture",
                    "replacement_text": "scalable microservices architecture on AWS",
                    "reason": "Add cloud context"
                }
            ]),
            
            # 6. Interview Questions
            json.dumps(["Describe your experience with microservices?"]),
            
            # 7. Cover Letter
            "Dear Hiring Manager,\n\nI am excited to apply...",
            
            # 8. Talking Points
            json.dumps(["I expanded the architecture to handle 2x traffic."])
        ])

        def side_effect(system, user, **kwargs):
            try:
                return next(responses_iter)
            except StopIteration:
                return ""

        with patch.object(analyzer, '_call_groq', side_effect=side_effect):
            result = analyzer.analyze(
                resume_text="Mock resume text", 
                job_title=self.job_title,
                job_description=self.job_description,
                success_profile=success_profile
            )
            
            self.assertEqual(result['scores']['skills'], 85)
            self.assertEqual(result['ats_score'], 92)
            self.assertEqual(len(result['suggestions']), 1)
            self.assertEqual(result['suggestions'][0]['original_text'], "microservices architecture")
            print("LLM Analyzer returned structured results successfully.")
            return result

    def test_4_resume_editing_preservation(self):
        """Test in-place editing preserves formatting."""
        print("\n--- Testing Resume Editor (Formatting Preservation) ---")
        
        editor = ResumeEditor(self.test_docx_path)
        
        # We want to replace "microservices architecture" which was created in the setup
        # It was: item1.add_run("microservices architecture").bold = True
        
        suggestions = [{
            "original_text": "microservices architecture",
            "replacement_text": "distributed cloud systems", # Change content
            "reason": "Testing replacement"
        }]
        
        result = editor.apply_suggestions(suggestions)
        
        self.assertEqual(result['applied'], 1, "Should have applied 1 suggestion")
        
        # Save and verify
        output_path = f"{self.output_dir}/test_optimized.docx"
        editor.save(output_path)
        print(f"Saved optimized resume to {output_path}")
        
        # Verification: Open the saved doc and check properties
        doc = Document(output_path)
        found_new_text = False
        is_bold = False
        
        for p in doc.paragraphs:
            for r in p.runs:
                if "distributed cloud systems" in r.text:
                    found_new_text = True
                    if r.bold:
                        is_bold = True
        
        self.assertTrue(found_new_text, "New text not found in document")
        self.assertTrue(is_bold, "Formatting lost! New text should be bold like the original.")
        print("Formatting preservation verified: Text replaced and Bold style kept.")

    def test_5_pdf_generation(self):
        """Test PDF generation for Interview Prep and Cover Letter."""
        print("\n--- Testing PDF Generation ---")
        
        gen = PDFGenerator()
        
        # Interview Prep
        questions = ["Q1: Explain Python GIL?", "Q2: How do you handle merge conflicts?"]
        int_path = f"{self.output_dir}/test_interview_prep.pdf"
        gen.generate_interview_prep(questions, self.job_title, int_path)
        self.assertTrue(os.path.exists(int_path))
        print(f"Generated {int_path}")
        
        # Cover Letter
        cl_text = "This is a test cover letter.\n\nIt has multiple paragraphs."
        cl_path = f"{self.output_dir}/test_cover_letter.pdf"
        gen.generate_cover_letter(cl_text, self.job_title, self.company_name, cl_path)
        self.assertTrue(os.path.exists(cl_path))
        print(f"Generated {cl_path}")

if __name__ == '__main__':
    unittest.main()
