"""
ResumeEditor – Reads, modifies, and saves .docx files while preserving formatting.

Uses python-docx for reading and lxml-level run manipulation for in-place
text replacement that keeps bold, italic, font, color, and all other
run-level formatting intact.
"""

import copy
import logging
import re
from typing import Optional

from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Inches
from lxml import etree

logger = logging.getLogger(__name__)


class ResumeEditor:
    """Read, analyze, and edit .docx files with formatting preservation."""

    def __init__(self, docx_path: str):
        """
        Load a .docx file for editing.

        Args:
            docx_path: Path to the .docx file
        """
        self.docx_path = docx_path
        self.document = Document(docx_path)
        logger.info("Loaded document: %s", docx_path)

    def extract_text(self) -> str:
        """
        Extract all text content from the document.

        Returns:
            Full text content with paragraphs separated by newlines
        """
        paragraphs = []
        for para in self.document.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # Also extract text from tables
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if text:
                            paragraphs.append(text)

        return "\n".join(paragraphs)

    def apply_suggestions(self, suggestions: list) -> dict:
        """
        Apply a list of text replacement suggestions to the document.

        Args:
            suggestions: List of dicts with 'original_text' and 'replacement_text'

        Returns:
            Dict with counts of applied and failed replacements
        """
        applied = 0
        failed = 0
        details = []

        for suggestion in suggestions:
            original = suggestion.get("original_text", "")
            replacement = suggestion.get("replacement_text", "")
            reason = suggestion.get("reason", "")

            if not original or not replacement:
                failed += 1
                details.append({"status": "skipped", "reason": "Empty text"})
                continue

            success = self._find_and_replace(original, replacement)
            if success:
                applied += 1
                details.append({
                    "status": "applied",
                    "original": original[:60] + "...",
                    "reason": reason,
                })
                logger.info("Applied: %s...", original[:50])
            else:
                failed += 1
                details.append({
                    "status": "failed",
                    "original": original[:60] + "...",
                    "reason": "Text not found in document",
                })
                logger.warning("Not found: %s...", original[:50])

        return {"applied": applied, "failed": failed, "details": details}

    def _find_and_replace(self, old_text: str, new_text: str) -> bool:
        """
        Find text across the document and replace it while preserving formatting.

        Searches all paragraphs (including those in tables) for the target text
        and performs a run-level replacement.
        """
        # Collect all paragraphs (body + tables)
        all_paragraphs = list(self.document.paragraphs)
        for table in self.document.tables:
            for row in table.rows:
                for cell in row.cells:
                    all_paragraphs.extend(cell.paragraphs)

        for paragraph in all_paragraphs:
            full_text = paragraph.text
            if old_text in full_text:
                self._replace_in_paragraph(paragraph, old_text, new_text)
                return True

        # Try normalized matching (collapse whitespace)
        normalized_old = re.sub(r'\s+', ' ', old_text).strip()
        for paragraph in all_paragraphs:
            normalized_para = re.sub(r'\s+', ' ', paragraph.text).strip()
            if normalized_old in normalized_para:
                # Use original paragraph text for replacement
                actual_old = self._find_actual_substring(paragraph.text, normalized_old)
                if actual_old:
                    self._replace_in_paragraph(paragraph, actual_old, new_text)
                    return True

        return False

    def _find_actual_substring(self, full_text: str, normalized_target: str) -> Optional[str]:
        """
        Find the actual substring in full_text that matches the normalized target.
        """
        # Slide a window across full_text
        words_target = normalized_target.split()
        if not words_target:
            return None

        # Find the start of the first word
        start_idx = 0
        while start_idx < len(full_text):
            pos = full_text.find(words_target[0], start_idx)
            if pos == -1:
                return None

            # Try to match all words from this position
            candidate_end = pos
            matched = True
            for word in words_target:
                # Skip whitespace
                while candidate_end < len(full_text) and full_text[candidate_end] == ' ':
                    candidate_end += 1
                if full_text[candidate_end:candidate_end + len(word)] == word:
                    candidate_end += len(word)
                else:
                    matched = False
                    break

            if matched:
                return full_text[pos:candidate_end]

            start_idx = pos + 1

        return None

    def _replace_in_paragraph(self, paragraph, old_text: str, new_text: str):
        """
        Replace text within a paragraph at the run level, preserving formatting.

        Strategy:
        1. Build a character-to-run map from concatenated run texts
        2. Find the span of old_text in the concatenated string
        3. Replace text in the affected runs:
           - First run: text before match + new_text
           - Middle runs: clear text (set to "")
           - Last run: text after match
        4. Run formatting XML (rPr) is NEVER touched
        """
        runs = paragraph.runs
        if not runs:
            return

        # Build concatenated text and character→run mapping
        concat = ""
        char_map = []  # char_map[i] = (run_index, char_offset_in_run)
        for run_idx, run in enumerate(runs):
            run_text = run.text or ""
            for char_offset, _ in enumerate(run_text):
                char_map.append((run_idx, char_offset))
            concat += run_text

        # Find old_text in concatenated string
        match_start = concat.find(old_text)
        if match_start == -1:
            return

        match_end = match_start + len(old_text)

        # Identify affected runs
        if match_start >= len(char_map) or match_end - 1 >= len(char_map):
            return

        first_run_idx, first_char_offset = char_map[match_start]
        last_run_idx, last_char_offset = char_map[match_end - 1]

        if first_run_idx == last_run_idx:
            # Simple case: old_text is entirely within one run
            run = runs[first_run_idx]
            run_text = run.text or ""
            end_offset = last_char_offset + 1
            run.text = run_text[:first_char_offset] + new_text + run_text[end_offset:]
        else:
            # Complex case: old_text spans multiple runs
            # First run: keep text before match, append new_text
            first_run = runs[first_run_idx]
            first_run_text = first_run.text or ""
            first_run.text = first_run_text[:first_char_offset] + new_text

            # Middle runs: clear their text
            for mid_idx in range(first_run_idx + 1, last_run_idx):
                runs[mid_idx].text = ""

            # Last run: keep text after match
            last_run = runs[last_run_idx]
            last_run_text = last_run.text or ""
            last_run.text = last_run_text[last_char_offset + 1:]

    def add_talking_points_page(self, talking_points: list):
        """
        Append a page break and a styled "Interview Talking Points" section.

        Args:
            talking_points: List of suggestion dicts with 'talking_point' fields
        """
        if not talking_points:
            return

        # Add page break
        self.document.add_page_break()

        # Add header
        heading = self.document.add_heading("Interview Talking Points", level=1)

        # Add subtitle
        subtitle = self.document.add_paragraph(
            "Prepared talking points to defend each resume optimization in interviews."
        )

        for i, tp in enumerate(talking_points):
            original = tp.get("original_text", "")
            replacement = tp.get("replacement_text", "")
            reason = tp.get("reason", "")
            point = tp.get("talking_point", "")

            # Section header
            self.document.add_heading(f"Edit {i + 1}: {tp.get('section', 'General')}", level=2)

            # Change description
            p = self.document.add_paragraph()
            runner = p.add_run("Changed: ")
            runner.bold = True
            p.add_run(f'"{original[:80]}..."' if len(original) > 80 else f'"{original}"')

            p2 = self.document.add_paragraph()
            runner2 = p2.add_run("To: ")
            runner2.bold = True
            p2.add_run(f'"{replacement[:80]}..."' if len(replacement) > 80 else f'"{replacement}"')

            if reason:
                p3 = self.document.add_paragraph()
                runner3 = p3.add_run("Why: ")
                runner3.bold = True
                p3.add_run(reason)

            if point:
                p4 = self.document.add_paragraph()
                runner4 = p4.add_run("Talking Point: ")
                runner4.bold = True
                runner4.italic = True
                p4.add_run(point)

            # Add spacing
            self.document.add_paragraph("")

    def save(self, output_path: str):
        """
        Save the modified document to a new path.

        Args:
            output_path: Path to save the modified .docx file
        """
        self.document.save(output_path)
        logger.info("Saved optimized document to: %s", output_path)
