
from docx import Document
from docx.shared import Pt, RGBColor

def create_rich_docx(filename):
    doc = Document()
    
    # Header
    h1 = doc.add_heading('John Doe', 0)
    h1.runs[0].font.color.rgb = RGBColor(0, 0, 0)
    
    p = doc.add_paragraph('San Francisco, CA | john.doe@example.com | (555) 123-4567')
    p.alignment = 1  # Center
    
    doc.add_heading('Summary', level=1)
    # A mix of formatting in one paragraph
    p = doc.add_paragraph()
    p.add_run("Experienced ").bold = True
    p.add_run("Backend Engineer with ")
    run = p.add_run("5+ years")
    run.italic = True
    run.font.color.rgb = RGBColor(255, 0, 0) # Red text to test color preservation
    p.add_run(" in building scalable systems using Python and Go.")
    
    doc.add_heading('Experience', level=1)
    
    # Job 1
    p = doc.add_paragraph()
    p.add_run("Senior Software Engineer").bold = True
    p.add_run(" | TechCorp Inc. | 2020 - Present")
    
    # Bullet points with formatting
    item1 = doc.add_paragraph(style='List Bullet')
    item1.add_run("Led the development of a ")
    item1.add_run("microservices architecture").bold = True
    item1.add_run(" that processed 1M+ requests per second.")
    
    item2 = doc.add_paragraph(style='List Bullet')
    item2.add_run("Reduced cloud costs by ")
    item2.add_run("30%").italic = True
    item2.add_run(" through optimization of AWS resources.")

    # Job 2
    p = doc.add_paragraph()
    p.add_run("Software Engineer").bold = True
    p.add_run(" | StartupXY | 2018 - 2020")
    
    item3 = doc.add_paragraph(style='List Bullet')
    item3.add_run("Implemented a ")
    item3.add_run("real-time chat feature").bold = True
    item3.add_run(" using WebSockets and Redis.")

    doc.save(filename)
    print(f"Created {filename}")

if __name__ == "__main__":
    create_rich_docx("tests/test_resume.docx")
